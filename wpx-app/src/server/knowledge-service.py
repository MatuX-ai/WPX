"""
Knowledge service — document upload, parsing, chunking, Chroma vector storage.

Endpoints:
  POST /api/knowledge/upload   — file or URL
  GET  /api/knowledge/list     — uploaded materials list
  GET  /api/knowledge/{id}/preview — full extracted text for preview

Run: python src/server/knowledge-service.py
"""

from __future__ import annotations

import json
import os
import re
import sys
import uuid
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import urlparse

try:
    from fastapi import FastAPI, File, Form, HTTPException, UploadFile
    from fastapi.middleware.cors import CORSMiddleware
    from fastapi.responses import JSONResponse
except ImportError:
    print("缺少 FastAPI，请运行：pip install -r src/server/requirements.txt", file=sys.stderr)
    raise

PORT = int(os.environ.get("KNOWLEDGE_SERVICE_PORT", "3003"))
MAX_UPLOAD_BYTES = int(os.environ.get("KNOWLEDGE_MAX_BYTES", str(20 * 1024 * 1024)))
CHUNK_SIZE = int(os.environ.get("KNOWLEDGE_CHUNK_SIZE", "800"))
CHUNK_OVERLAP = int(os.environ.get("KNOWLEDGE_CHUNK_OVERLAP", "100"))

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data" / "knowledge"
DOCS_DIR = DATA_DIR / "docs"
MANIFEST_PATH = DATA_DIR / "manifest.json"
CHROMA_DIR = DATA_DIR / "chroma"

ALLOWED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "word",
    ".md": "markdown",
    ".markdown": "markdown",
    ".txt": "text",
}

app = FastAPI(title="WPX Knowledge Service")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

_chroma_collection = None
_embedding_backend = "unknown"


class HashFallbackEmbeddingFunction:
    """Deterministic lightweight embedding for dev/acceptance when ML models are unavailable."""

    @staticmethod
    def name() -> str:
        return "hash-fallback"

    def default_space(self) -> str:
        return "cosine"

    def supported_spaces(self) -> list[str]:
        return ["cosine", "l2", "ip"]

    @staticmethod
    def build_from_config(config: dict[str, Any]) -> "HashFallbackEmbeddingFunction":
        return HashFallbackEmbeddingFunction()

    def get_config(self) -> dict[str, Any]:
        return {}

    def __call__(self, input):  # noqa: A002
        import hashlib
        import math

        vectors = []
        for text in input:
            digest = hashlib.sha256(str(text).encode("utf-8")).digest()
            values = []
            for index in range(384):
                byte = digest[index % len(digest)]
                values.append((byte / 255.0) * 2 - 1)
            norm = math.sqrt(sum(value * value for value in values)) or 1.0
            vectors.append([value / norm for value in values])
        return vectors


def create_embedding_function():
    global _embedding_backend

    if os.environ.get("KNOWLEDGE_SKIP_ONNX", "").lower() in ("1", "true", "yes"):
        print("[knowledge-service] 已设置 KNOWLEDGE_SKIP_ONNX，跳过 ONNX 模型下载")
    else:
        try:
            from chromadb.utils.embedding_functions import DefaultEmbeddingFunction

            embedding_fn = DefaultEmbeddingFunction()
            _embedding_backend = "chromadb-default"
            return embedding_fn
        except Exception as exc:
            print(
                "[knowledge-service] DefaultEmbeddingFunction 不可用，尝试 sentence-transformers:",
                exc,
            )

    if os.environ.get("KNOWLEDGE_USE_SENTENCE_TRANSFORMERS", "").lower() in ("1", "true", "yes"):
        try:
            from chromadb.utils.embedding_functions import SentenceTransformerEmbeddingFunction

            embedding_fn = SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")
            _embedding_backend = "sentence-transformers"
            return embedding_fn
        except Exception as exc:
            print(
                "[knowledge-service] sentence-transformers 不可用，回退到 hash embedding:",
                exc,
            )

    _embedding_backend = "hash-fallback"
    return HashFallbackEmbeddingFunction()


def get_chroma_collection():
    global _chroma_collection
    if _chroma_collection is not None:
        return _chroma_collection

    try:
        import chromadb
    except ImportError as exc:
        raise RuntimeError(
            "缺少 chromadb，请运行：pip install -r src/server/requirements.txt"
        ) from exc

    embedding_fn = create_embedding_function()
    collection_name = f"knowledge-{embedding_fn.name()}"

    client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    _chroma_collection = client.get_or_create_collection(
        name=collection_name,
        embedding_function=embedding_fn,
        metadata={"hnsw:space": "cosine"},
    )
    return _chroma_collection


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def ensure_dirs() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    CHROMA_DIR.mkdir(parents=True, exist_ok=True)


def load_manifest() -> list[dict[str, Any]]:
    ensure_dirs()
    if not MANIFEST_PATH.exists():
        return []
    try:
        data = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))
        return data if isinstance(data, list) else []
    except (json.JSONDecodeError, OSError):
        return []


def save_manifest(items: list[dict[str, Any]]) -> None:
    ensure_dirs()
    MANIFEST_PATH.write_text(
        json.dumps(items, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("缺少 pypdf") from exc

    reader = PdfReader(BytesIO(data))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return normalize_text("\n".join(parts))


def extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("缺少 python-docx") from exc

    doc = Document(BytesIO(data))
    parts = [para.text for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return normalize_text("\n".join(parts))


def extract_text_file(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk", "latin-1"):
        try:
            return normalize_text(data.decode(encoding))
        except UnicodeDecodeError:
            continue
    return normalize_text(data.decode("utf-8", errors="replace"))


def extract_url(url: str) -> tuple[str, str]:
    parsed = urlparse(url.strip())
    if parsed.scheme not in ("http", "https") or not parsed.netloc:
        raise ValueError("请输入有效的 http/https URL")

    try:
        import trafilatura
    except ImportError as exc:
        raise RuntimeError("缺少 trafilatura") from exc

    downloaded = trafilatura.fetch_url(url)
    if not downloaded:
        raise ValueError("无法抓取该网页，请检查 URL 是否可访问")

    text = trafilatura.extract(downloaded, include_comments=False, include_tables=True)
    if not text or not text.strip():
        raise ValueError("未能从网页提取正文内容")

    title = trafilatura.extract_metadata(downloaded)
    filename = (title.title if title and title.title else parsed.netloc) + ".web"
    return filename, normalize_text(text)


def parse_upload(filename: str, data: bytes) -> tuple[str, str]:
    ext = Path(filename).suffix.lower()
    doc_type = ALLOWED_EXTENSIONS.get(ext)
    if not doc_type:
        allowed = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise ValueError(f"不支持的文件类型，仅支持：{allowed}")

    if doc_type == "pdf":
        content = extract_pdf(data)
    elif doc_type == "word":
        content = extract_docx(data)
    else:
        content = extract_text_file(data)

    if not content:
        raise ValueError("未能从文件中提取文本内容")

    return doc_type, content


def chunk_text(text: str) -> list[str]:
    if len(text) <= CHUNK_SIZE:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + CHUNK_SIZE, len(text))
        if end < len(text):
            boundary = text.rfind("\n", start, end)
            if boundary > start + CHUNK_SIZE // 2:
                end = boundary + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = max(end - CHUNK_OVERLAP, start + 1)

    return chunks or [text]


def store_document(
    filename: str,
    doc_type: str,
    content: str,
    source: str = "file",
) -> dict[str, Any]:
    doc_id = str(uuid.uuid4())
    uploaded_at = utc_now_iso()

    text_path = DOCS_DIR / f"{doc_id}.txt"
    text_path.write_text(content, encoding="utf-8")

    chunks = chunk_text(content)
    collection = get_chroma_collection()

    ids = [f"{doc_id}_{i}" for i in range(len(chunks))]
    metadatas = [
        {
            "doc_id": doc_id,
            "filename": filename,
            "type": doc_type,
            "source": source,
            "chunk_index": i,
            "uploaded_at": uploaded_at,
        }
        for i in range(len(chunks))
    ]

    collection.add(ids=ids, documents=chunks, metadatas=metadatas)

    item = {
        "id": doc_id,
        "filename": filename,
        "type": doc_type,
        "source": source,
        "uploadedAt": uploaded_at,
        "charCount": len(content),
        "chunkCount": len(chunks),
    }

    manifest = load_manifest()
    manifest.insert(0, item)
    save_manifest(manifest)

    return item


def find_document(doc_id: str) -> dict[str, Any] | None:
    for item in load_manifest():
        if item.get("id") == doc_id:
            return item
    return None


def clear_knowledge_index() -> None:
    """清空上传资料索引与本地缓存，不触及用户磁盘上的原始文件。"""
    save_manifest([])

    if DOCS_DIR.exists():
        for path in DOCS_DIR.glob("*.txt"):
            try:
                path.unlink(missing_ok=True)
            except OSError:
                pass

    global _chroma_collection
    try:
        collection = get_chroma_collection()
        existing = collection.get()
        ids = existing.get("ids") if isinstance(existing, dict) else None
        if ids:
            collection.delete(ids=ids)
    except Exception:
        _chroma_collection = None


@app.get("/api/knowledge/health")
async def health():
    chroma_ok = False
    try:
        get_chroma_collection()
        chroma_ok = True
    except Exception:
        chroma_ok = False

    return {
        "status": "ok",
        "chroma": chroma_ok,
        "embedding": _embedding_backend,
        "documents": len(load_manifest()),
    }


@app.get("/api/knowledge/list")
async def list_documents():
    items = load_manifest()
    return {"items": items}


@app.get("/api/knowledge/{doc_id}/preview")
async def preview_document(doc_id: str):
    item = find_document(doc_id)
    if not item:
        raise HTTPException(status_code=404, detail="资料不存在")

    text_path = DOCS_DIR / f"{doc_id}.txt"
    if not text_path.exists():
        raise HTTPException(status_code=404, detail="资料内容不存在")

    content = text_path.read_text(encoding="utf-8")
    return {
        "id": item["id"],
        "filename": item["filename"],
        "type": item["type"],
        "uploadedAt": item["uploadedAt"],
        "content": content,
    }


@app.post("/api/knowledge/clear")
async def clear_knowledge_cache():
    clear_knowledge_index()
    return {"success": True}


@app.post("/api/knowledge/search")
async def search_knowledge(
    query: str = Form(...),
    top_k: int = Form(5),
):
    """
    语义搜索知识库。

    请求：POST /api/knowledge/search
      - query (str): 搜索查询文本
      - top_k (int): 返回的最相关片段数（默认 5，最大 20）

    响应：
      {
        "results": [
          {
            "chunk_id": "xxx_0",
            "doc_id": "xxx",
            "filename": "xxx.pdf",
            "text": "匹配的文本片段...",
            "score": 0.85
          },
          ...
        ],
        "query": "原始查询",
        "backend": "sentence-transformers",
        "total_indexed": 42
      }
    """
    if not query or not query.strip():
        raise HTTPException(status_code=400, detail="query 不能为空")

    k = max(1, min(20, int(top_k)))

    try:
        collection = get_chroma_collection()
        results = collection.query(
            query_texts=[query.strip()],
            n_results=k,
            include=["documents", "metadatas", "distances"],
        )
    except Exception as exc:
        raise HTTPException(
            status_code=500, detail=f"向量检索失败：{exc}"
        ) from exc

    if not results or not results.get("ids") or not results["ids"][0]:
        return {
            "results": [],
            "query": query.strip(),
            "backend": _embedding_backend,
            "total_indexed": len(load_manifest()),
        }

    items = []
    ids = results["ids"][0]
    docs = results.get("documents", [[]])[0]
    metas = results.get("metadatas", [[]])[0]
    dists = results.get("distances", [[]])[0]

    for i in range(len(ids)):
        meta = metas[i] if i < len(metas) else {}
        items.append({
            "chunk_id": ids[i],
            "doc_id": meta.get("doc_id", ""),
            "filename": meta.get("filename", ""),
            "type": meta.get("type", ""),
            "text": docs[i] if i < len(docs) else "",
            "score": round(1.0 - dists[i], 4) if i < len(dists) and dists[i] is not None else None,
        })

    return {
        "results": items,
        "query": query.strip(),
        "backend": _embedding_backend,
        "total_indexed": len(load_manifest()),
    }


@app.post("/api/knowledge/upload")
async def upload_document(
    file: UploadFile | None = File(None),
    url: str | None = Form(None),
):
    try:
        if file and file.filename:
            data = await file.read()
            if not data:
                raise HTTPException(status_code=400, detail="文件为空")
            if len(data) > MAX_UPLOAD_BYTES:
                raise HTTPException(
                    status_code=413,
                    detail=f"文件过大，最大支持 {MAX_UPLOAD_BYTES // (1024 * 1024)}MB",
                )

            doc_type, content = parse_upload(file.filename, data)
            item = store_document(file.filename, doc_type, content, source="file")
            return {"success": True, "item": item}

        if url and url.strip():
            filename, content = extract_url(url.strip())
            item = store_document(filename, "web", content, source="url")
            return {"success": True, "item": item}

        raise HTTPException(status_code=400, detail="请上传文件或提供 URL")
    except HTTPException:
        raise
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"处理失败：{exc}") from exc


@app.exception_handler(HTTPException)
async def http_exception_handler(_request, exc: HTTPException):
    detail = exc.detail
    if isinstance(detail, list):
        message = "; ".join(str(item) for item in detail)
    else:
        message = str(detail)

    return JSONResponse(
        status_code=exc.status_code,
        content={"error": message, "detail": detail},
    )


# =========================
# 标签提取与自动索引
# =========================

# 中文常见停用词
_STOPWORDS = set(
    "的 了 在 是 我 有 和 就 不 人 都 一 一个 上 也 很 到 说 要 去 你 "
    "会 着 没有 看 好 自己 这 他 她 它 们 那 些 什么 而 为 所以 因为 "
    "但是 如果 虽然 可以 这个 那个 这些 那些 我们 他们 她们 它们 "
    "已经 知道 觉得 认为 应该 能够 需要 对于 关于 通过 根据 按照 "
    "之后 以前 以后 现在 当时 目前 最近 主要 一般 比较 非常 "
    "还是 只是 然后 不过 以及 并且 或者 此外 另外 比如 例如 "
    "其中 其他 所有 每个 整个 部分 方面 问题 情况 作用 结果 "
    "关系 过程 方法 方式 内容 形式 目的 意义 原因 条件 基础 "
    "发展 形成 进行 使用 利用 具有 存在 出现 产生 发生 成为 "
    "不是 没有 不能 不会 可能 是否 是否 更加 越来越 等等".split()
)


def extract_tags(content: str, top_n: int = 10) -> list[dict[str, Any]]:
    """从文档内容中提取关键词标签（基于词频）"""
    if not content:
        return []

    text = re.sub("[，。！？、；：\u201c\u201d\u2018\u2019（）\\[\\]【】《》\\s]+", " ", content)

    # 提取中文 2-4 字词和英文 3+ 字母词
    cn_words = re.findall(r"[\u4e00-\u9fff]{2,4}", text)
    en_words = re.findall(r"[a-zA-Z]{3,}", text.lower())

    freq: dict[str, int] = {}
    for w in cn_words:
        if w not in _STOPWORDS and len(w) >= 2:
            freq[w] = freq.get(w, 0) + 1
    for w in en_words:
        if w not in _STOPWORDS:
            freq[w] = freq.get(w, 0) + 1

    # 按词频排序
    sorted_tags = sorted(freq.items(), key=lambda x: -x[1])[:top_n]
    max_freq = sorted_tags[0][1] if sorted_tags else 1

    return [
        {"tag": word, "count": count, "weight": round(count / max_freq, 2)}
        for word, count in sorted_tags
    ]


def detect_cross_references(
    current_doc_id: str,
    content: str,
    all_items: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """检测文档中引用的其他文档（通过文件名提及检测）"""
    refs = []
    for item in all_items:
        if item["id"] == current_doc_id:
            continue
        filename = item.get("filename", "")
        if not filename:
            continue
        # 去掉扩展名的文件名
        basename = Path(filename).stem
        if len(basename) >= 2 and basename in content:
            refs.append({
                "doc_id": item["id"],
                "filename": filename,
                "type": item.get("type", ""),
                "mentions": content.count(basename),
            })
    return sorted(refs, key=lambda x: -x["mentions"])


@app.get("/api/knowledge/tags")
async def get_tag_cloud():
    """
    获取标签云 — 所有文档的关键词聚合
    响应: { tags: [{ tag, count, weight, docCount }], totalDocs: int }
    """
    items = load_manifest()
    all_tags: dict[str, dict[str, Any]] = {}

    for item in items:
        doc_path = DOCS_DIR / f"{item['id']}.txt"
        if not doc_path.exists():
            continue
        try:
            content = doc_path.read_text(encoding="utf-8")
        except Exception:
            continue

        doc_tags = extract_tags(content, top_n=10)
        for t in doc_tags:
            tag = t["tag"]
            if tag not in all_tags:
                all_tags[tag] = {"tag": tag, "count": 0, "weight": 0, "docCount": 0}
            all_tags[tag]["count"] += t["count"]
            all_tags[tag]["weight"] = max(all_tags[tag]["weight"], t["weight"])
            all_tags[tag]["docCount"] += 1

    sorted_tags = sorted(all_tags.values(), key=lambda x: -x["count"])[:50]
    return {"tags": sorted_tags, "totalDocs": len(items)}


@app.get("/api/knowledge/index")
async def get_auto_index():
    """
    自动索引页 — 按标签分组显示文档
    响应: { groups: [{ tag, docs: [{ id, filename, type, uploadedAt, tags }] }] }
    """
    items = load_manifest()
    tag_groups: dict[str, list[dict[str, Any]]] = {}

    for item in items:
        doc_path = DOCS_DIR / f"{item['id']}.txt"
        if not doc_path.exists():
            continue
        try:
            content = doc_path.read_text(encoding="utf-8")
        except Exception:
            continue

        doc_tags = extract_tags(content, top_n=5)
        doc_info = {
            "id": item["id"],
            "filename": item["filename"],
            "type": item.get("type", ""),
            "uploadedAt": item["uploadedAt"],
            "charCount": len(content),
            "tags": [t["tag"] for t in doc_tags[:3]],
        }

        for t in doc_tags[:3]:
            tag = t["tag"]
            if tag not in tag_groups:
                tag_groups[tag] = []
            # 避免同一文档在同一标签下重复
            if not any(d["id"] == doc_info["id"] for d in tag_groups[tag]):
                tag_groups[tag].append(doc_info)

    # 按组内文档数量排序
    groups = [
        {"tag": tag, "docs": docs}
        for tag, docs in sorted(tag_groups.items(), key=lambda x: -len(x[1]))[:20]
    ]
    return {"groups": groups, "totalDocs": len(items)}


@app.get("/api/knowledge/{doc_id}/links")
async def get_document_links(doc_id: str):
    """
    获取文档的跨文档引用关系
    响应: { docId, references: [{ doc_id, filename, type, mentions }] }
    """
    items = load_manifest()
    item = next((i for i in items if i["id"] == doc_id), None)
    if not item:
        raise HTTPException(status_code=404, detail="资料不存在")

    doc_path = DOCS_DIR / f"{doc_id}.txt"
    if not doc_path.exists():
        return {"docId": doc_id, "references": []}

    content = doc_path.read_text(encoding="utf-8")
    refs = detect_cross_references(doc_id, content, items)
    return {"docId": doc_id, "filename": item["filename"], "references": refs}


@app.post("/api/knowledge/fulltext-search")
async def fulltext_search(
    query: str = Form(...),
    limit: int = Form(20),
):
    """
    全文搜索 — 在所有文档中进行简单的文本匹配（非语义）
    用于补充语义搜索，在精确关键词匹配场景更有效

    请求: POST /api/knowledge/fulltext-search
      - query (str): 搜索关键词
      - limit (int): 最大返回数
    """
    if not query or not query.strip():
        raise HTTPException(status_code=400, detail="query 不能为空")

    q = query.strip()
    items = load_manifest()
    results = []

    for item in items:
        doc_path = DOCS_DIR / f"{item['id']}.txt"
        if not doc_path.exists():
            continue
        try:
            content = doc_path.read_text(encoding="utf-8")
        except Exception:
            continue

        # 简单的关键词匹配分数
        score = content.lower().count(q.lower())
        if score == 0:
            # 分词匹配
            words = q.split()
            score = sum(content.lower().count(w.lower()) for w in words)

        if score > 0:
            # 提取匹配片段
            idx = content.lower().find(q.lower())
            if idx < 0 and len(q.split()) > 1:
                idx = content.lower().find(q.split()[0].lower())
            snippet_start = max(0, idx - 100) if idx >= 0 else 0
            snippet = content[snippet_start:snippet_start + 300]

            results.append({
                "doc_id": item["id"],
                "filename": item["filename"],
                "type": item.get("type", ""),
                "uploadedAt": item["uploadedAt"],
                "charCount": len(content),
                "matchCount": score,
                "snippet": snippet,
            })

    results.sort(key=lambda x: -x["matchCount"])
    return {
        "results": results[:limit],
        "query": q,
        "totalHits": len(results),
        "totalIndexed": len(items),
    }


if __name__ == "__main__":
    import uvicorn

    ensure_dirs()
    print(f"[knowledge-service] 运行于 http://localhost:{PORT}")
    print("[knowledge-service] POST /api/knowledge/upload")
    print("[knowledge-service] GET  /api/knowledge/list")
    print("[knowledge-service] GET  /api/knowledge/{id}/preview")

    uvicorn.run(app, host="0.0.0.0", port=PORT)
